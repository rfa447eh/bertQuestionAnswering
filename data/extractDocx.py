from docx import Document

# Function to read a .docx file and return the text
def read_docx(file_path):
    # Open the Word document
    doc = Document(file_path)

    # Initialize an empty list to hold paragraphs
    paragraphs = []

    # Loop through each paragraph in the document and extract the text
    for paragraph in doc.paragraphs:
        paragraphs.append(paragraph.text)

    return paragraphs

# Function to process paragraphs and extract specific ones
def process_paragraphs(paragraphs):
    specific_paragraphs = []
    remaining_paragraphs = []

    for para in paragraphs:
        if para.startswith("What") or para.startswith("How"):
            specific_paragraphs.append(para)
            remaining_paragraphs.append(para)
        else:
            remaining_paragraphs.append(para)

    return specific_paragraphs, remaining_paragraphs

# Path to the .docx file
file_path = './MatterOverview.docx'

# Read the content of the .docx file
paragraphs = read_docx(file_path)

# Process paragraphs to extract specific ones and get remaining ones
specific_paragraphs, remaining_paragraphs = process_paragraphs(paragraphs)

# Convert lists of paragraphs to strings
specific_paragraphs_text = '\n'.join(specific_paragraphs)
remaining_paragraphs_text = '\n'.join(remaining_paragraphs)

# Print the results
print("Specific Paragraphs:")
print(specific_paragraphs_text)
print("\nRemaining Paragraphs:")
print(remaining_paragraphs_text)

# Optional: Save the specific paragraphs to a new .docx file
from docx import Document as DocxDocument

def save_paragraphs_to_docx(paragraphs, file_path):
    doc = DocxDocument()
    for para in paragraphs:
        doc.add_paragraph(para)
    doc.save(file_path)

# Path for saving the specific paragraphs
output_file_path = './SpecificParagraphs.docx'
save_paragraphs_to_docx(specific_paragraphs, output_file_path)


